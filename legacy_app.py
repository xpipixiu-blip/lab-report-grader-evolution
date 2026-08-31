"""
实验报告 AI 批改系统
==============================
两轮批改流程：
  第一轮：AI 用内置标准快速预批改 → 浏览器查看结果
  第二轮：教师调整评分侧重 → AI 优化客制化规则 → 重新批改 → 生成 Excel
"""
import gradio as gr
import os, re, json, tempfile, zipfile, copy, base64, io, html
import urllib.request, urllib.error
from pathlib import Path
from datetime import datetime
from typing import Generator, Optional
import yaml
from openai import OpenAI
from docx import Document
import fitz
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from grader_core.injection_guard import SECURITY_RULES, protect_report_text, scan_report

CONFIG_PATH = Path(__file__).parent / "config.yaml"

def load_config() -> dict:
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

# ============================================================
# 1. 文件文本提取（不变）
# ============================================================
def extract_text_from_docx(filepath: str) -> str:
    doc = Document(filepath)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)

def extract_text_from_pdf(filepath: str) -> str:
    doc = fitz.open(filepath)
    parts = []
    for page in doc:
        t = page.get_text("text")
        if t.strip():
            parts.append(t.strip())
    doc.close()
    return "\n\n".join(parts)

def extract_text(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext == ".docx": return extract_text_from_docx(filepath)
    elif ext == ".pdf": return extract_text_from_pdf(filepath)
    else: raise ValueError(f"不支持的文件格式: {ext}")

# ============================================================
# 2. 学生信息提取（不变）
# ============================================================
def extract_student_info(filename: str) -> dict:
    basename = Path(filename).stem
    info = {"student_id": "", "student_name": ""}
    m = re.search(r'(\d{10,14})\s*([一-鿿]{2,3})', basename)
    if m:
        info["student_id"] = m.group(1)
        info["student_name"] = m.group(2)
        return info
    m = re.search(r'(\d{10,14})', basename)
    if m: info["student_id"] = m.group(1)
    m = re.search(r'([一-鿿]{2,3})', basename)
    if m: info["student_name"] = m.group(1)
    if info["student_id"] or info["student_name"]:
        return info
    info["student_name"] = basename[:30]
    return info

# ============================================================
# 2.5 图片提取（多模态支持）
# ============================================================
MAX_IMAGES = 8  # 每份报告最多发给 AI 的图片数

def extract_images_from_docx(filepath: str) -> list[str]:
    """从 .docx 中提取嵌入图片，返回 base64 data URI 列表"""
    images = []
    try:
        doc = Document(filepath)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_part = rel.target_part
                    img_bytes = img_part.blob
                    if len(img_bytes) < 512:  # 跳过太小的图（可能是图标/装饰）
                        continue
                    ext = (img_part.ext or "png").lstrip(".")
                    mime = f"image/{ext}" if ext in ("png","jpeg","jpg","gif","webp","bmp") else "image/png"
                    b64 = base64.b64encode(img_bytes).decode("ascii")
                    images.append(f"data:{mime};base64,{b64}")
                    if len(images) >= MAX_IMAGES:
                        break
                except Exception:
                    continue
    except Exception:
        pass
    return images

def extract_images_from_pdf(filepath: str) -> list[str]:
    """从 .pdf 中提取嵌入图片，返回 base64 data URI 列表"""
    images = []
    try:
        doc = fitz.open(filepath)
        for page in doc:
            img_list = page.get_images()
            for img in img_list:
                try:
                    xref = img[0]
                    base_img = doc.extract_image(xref)
                    img_bytes = base_img["image"]
                    if len(img_bytes) < 512:
                        continue
                    ext = base_img.get("ext", "png")
                    mime = f"image/{ext}" if ext in ("png","jpeg","jpg","gif","webp","bmp") else "image/png"
                    b64 = base64.b64encode(img_bytes).decode("ascii")
                    images.append(f"data:{mime};base64,{b64}")
                    if len(images) >= MAX_IMAGES:
                        break
                except Exception:
                    continue
            if len(images) >= MAX_IMAGES:
                break
        doc.close()
    except Exception:
        pass
    return images

def extract_images(filepath: str) -> list[str]:
    """根据文件类型提取嵌入图片"""
    ext = Path(filepath).suffix.lower()
    if ext == ".docx": return extract_images_from_docx(filepath)
    elif ext == ".pdf": return extract_images_from_pdf(filepath)
    return []

# ============================================================
# 3. 模型列表获取（不变）
# ============================================================
def fetch_model_list(api_key: str, base_url: str) -> list[str]:
    if not api_key or not base_url: return []
    base = base_url.rstrip("/")
    if base.endswith("/chat/completions"): base = base.rsplit("/chat/completions", 1)[0]
    try:
        req = urllib.request.Request(f"{base}/models",
            headers={"Authorization": f"Bearer {api_key}","Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        if "data" in data:
            models = [m["id"] for m in data["data"] if not any(
                kw in m["id"].lower() for kw in ["embed","moderat","whisper","tts","dall-e","vision-preview"])]
            return sorted(models)
    except Exception: pass
    return []

def get_effective_base_url(raw: str) -> str:
    url = raw.strip().rstrip("/")
    if url.endswith("/chat/completions"): url = url.rsplit("/chat/completions", 1)[0]
    return url

# ============================================================
# 4. AI 评分核心
# ============================================================
def build_system_prompt(rubric: dict, experiment_name: str, extra_instructions: str = "") -> str:
    """构建评分 system prompt，支持额外的客制化指令"""
    report_type = rubric["report_type"]
    total_score = rubric["total_score"]
    dim_lines, dim_keys = [], []
    for i, dim in enumerate(rubric["dimensions"], 1):
        dim_lines.append(
            f"### 维度{i}：{dim['name']}（满分 {dim['max_score']} 分，权重 {int(dim['weight']*100)}%）\n"
            f"{dim['description'].strip()}")
        dim_keys.append(dim["key"])
    dim_text = "\n\n".join(dim_lines)

    # 客制化指令
    custom_block = ""
    if extra_instructions:
        custom_block = f"""
## ⚠️ 本次批改特别要求（优先级高于上述通用标准）

{extra_instructions}

"""

    prompt = f"""{SECURITY_RULES}

你是一位严格但富有建设性的大学无机化学实验课助教，正在批改学生的**{report_type}**。{f"实验名称：{experiment_name}" if experiment_name else ""}

## 评分标准（满分 {total_score} 分）

{dim_text}
{custom_block}
## ⚠️ 技术性问题豁免（最高优先级）

**以下情况一律视为系统提取/显示故障，严禁因此扣分：**
- 化学式缺失或显示为乱码/空白（如本该写 H₂SO₄ 的地方为空或乱码）
- 试剂名称缺失或显示异常
- 化学反应方程式未配平、箭头方向错误等格式问题
- 下标、上标丢失（如 H2O 而非 H₂O）

**你的评分应聚焦于：** 报告内容的完整性、逻辑性、学生对实验流程和原理的理解深度，而非文本格式或化学式显示是否正确。如果在评语中提到上述问题，只能作为善意提醒（如「此处化学式可能因格式转换未能正确显示，建议检查原文」），绝不能作为扣分理由。

## 评分要求

1. **逐维度评分**：认真阅读学生报告，在每个维度上根据标准的描述给出分数和简短具体评语
2. **总分**：各维度分数相加（带小数的保留1位）
3. **综合评语**：50-150字，先肯定优点，再指出主要不足和具体改进建议，语言要鼓励性但专业
4. **严格区分度**：优秀（90%+）、良好（80-89%）、中等（70-79%）、及格（60-69%）、不及格（<60%）

## 输出格式

你必须严格按照以下 JSON 格式输出（不要加代码块标记，直接输出纯 JSON）：

{{
  "dimensions": [
    {{"key": "{dim_keys[0]}", "score": 0.0, "comment": "对该维度的简短评语"}},
    ...
  ],
  "total_score": 0.0,
  "overall_comment": "综合评语"
}}

注意：
- 维度顺序必须与上面列出的顺序一致
- score 必须是数字（不是字符串），不超过该维度满分
- overall_comment 必须是一段完整的话，不要分点
- 评分要严格，不要给虚高的分数"""
    return prompt

def build_user_prompt(report_text: str, student_name: str, student_id: str,
                      report_type: str, experiment_name: str) -> str:
    max_chars = 30000
    if len(report_text) > max_chars:
        report_text = report_text[:max_chars] + "\n\n... [报告过长，已截断]"
    header = f"**学生**：{student_name}（{student_id}）"
    if experiment_name: header += f"\n**实验名称**：{experiment_name}"
    header += f"\n**报告类型**：{report_type}"
    protected_report, _ = protect_report_text(report_text)
    return f"""{header}

以下为学生提交的报告内容：

---
{protected_report}
---"""

def parse_grading_response(response_text: str, rubric: dict) -> Optional[dict]:
    cleaned = response_text.strip()
    try: return json.loads(cleaned)
    except json.JSONDecodeError: pass
    m = re.search(r'```(?:json)?\s*\n?([\s\S]*?)\n?```', cleaned)
    if m:
        try: return json.loads(m.group(1).strip())
        except json.JSONDecodeError: pass
    m = re.search(r'\{[\s\S]*\}', cleaned)
    if m:
        try: return json.loads(m.group(0))
        except json.JSONDecodeError: pass
    try:
        fixed = cleaned.replace('“','"').replace('”','"').replace('‘',"'").replace('’',"'")
        start, end = fixed.find('{'), fixed.rfind('}')
        if start >= 0 and end > start: return json.loads(fixed[start:end+1])
    except json.JSONDecodeError: pass
    return None

def grade_one_report(client, model, report_text, student_info, rubric, experiment_name, config,
                     extra_instructions="", images=None):
    """批改一份报告。如果提供 images（base64 URI 列表），使用多模态格式"""
    system_prompt = build_system_prompt(rubric, experiment_name, extra_instructions)

    # 如果有图片，在 system prompt 末尾追加提示
    if images:
        system_prompt += ("\n\n**注意：学生报告中包含图片（化学方程式、图表、曲线图等），"
                          "请仔细查看图片中的内容（包括图标题、坐标轴标注、数据点等），"
                          "结合图片和文字内容综合评分。图片中的公式/方程式请勿视为乱码。**")

    user_text = build_user_prompt(report_text,
        student_info.get("student_name",""), student_info.get("student_id",""),
        rubric["report_type"], experiment_name)

    # 构建 user message：有图则用多模态数组，无图则纯文本
    if images:
        user_content = [{"type": "text", "text": user_text + "\n\n**以下为报告中嵌入的图片：**"}]
        for img in images:
            user_content.append({"type": "image_url", "image_url": {"url": img, "detail": "auto"}})
    else:
        user_content = user_text

    api_cfg = config.get("api", {})

    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content},
                ],
                temperature=api_cfg.get("temperature",0.3),
                max_tokens=api_cfg.get("max_tokens",4096),
                timeout=api_cfg.get("timeout",120))
            content = response.choices[0].message.content
            if not content:
                if attempt < 2: continue
                return {"error":True,"message":"AI 返回了空内容","student_info":student_info}
            parsed = parse_grading_response(content, rubric)
            if parsed:
                if "dimensions" in parsed and "total_score" in parsed:
                    calc_total = sum(d.get("score",0) for d in parsed["dimensions"])
                    if abs(calc_total - parsed["total_score"]) > 0.5:
                        parsed["total_score"] = round(calc_total, 1)
                return parsed
            else:
                if attempt < 2: continue
                return {"error":True,"message":f"无法解析 AI 返回的 JSON。原始返回：\n{content[:500]}",
                        "student_info":student_info,"raw_response":content[:1000]}
        except Exception as e:
            # 如果多模态失败且是第一次尝试，可能是模型不支持图片，降级为纯文本
            err_msg = str(e)
            if images and attempt == 0 and ("image" in err_msg.lower() or "multipart" in err_msg.lower() or "vision" in err_msg.lower()):
                # 降级：去掉图片重试
                images = None
                user_content = user_text
                continue
            if attempt < 2: continue
            return {"error":True,"message":f"API 调用失败：{err_msg}","student_info":student_info}
    return {"error":True,"message":"未知错误","student_info":student_info}

# ============================================================
# 5. 客制化评分规则生成
# ============================================================
def generate_customized_rubric(client, model, rubric, adjustment_notes, special_situations,
                                first_results, experiment_name, config):
    """让 AI 根据教师的调整意见，在内置标准基础上生成增量修改指令"""

    # 构建完整的当前评分标准（不是摘要，是完整内容）
    rubric_full = []
    for i, d in enumerate(rubric["dimensions"], 1):
        rubric_full.append(
            f"### 维度{i}：{d['name']}（满分 {d['max_score']} 分，权重 {int(d['weight']*100)}%）\n"
            f"{d['description'].strip()}")
    rubric_text = "\n\n".join(rubric_full)

    total_score = rubric["total_score"]
    report_type = rubric["report_type"]

    # 汇总第一轮结果
    valid = [r for r in first_results if not r.get("error")]
    scores = [r["total_score"] for r in valid]
    avg_s = sum(scores)/len(scores) if scores else 0

    prompt = f"""你是一位无机化学实验课的资深教师。你的任务是：根据另一位教师的反馈意见，在**现有完整评分标准的基础上**，生成一份「增量修改说明」。

## ⚠️ 重要：以下是已有的完整评分标准，必须作为基线保留

{rubric_text}

## 第一轮预批改概况
- 报告类型：{report_type}（满分 {total_score} 分）
- 有效批改数：{len(valid)} 份
- 平均分：{round(avg_s,2)}/{total_score}，最高 {max(scores) if scores else 'N/A'}，最低 {min(scores) if scores else 'N/A'}
{f"实验名称：{experiment_name}" if experiment_name else ""}

## 教师的修改意见

**需要调整的方面：**
{adjustment_notes if adjustment_notes else "（无——教师对现有标准满意，只需处理特殊情况）"}

**特殊情况（个别学生的加分/免扣分等）：**
{special_situations if special_situations else "（无特殊情况）"}

## 你的任务：生成增量修改指令

**核心原则：只改教师提到的，没提到的一律保留原标准。**

你需要生成一段将直接插入到批改 AI 的 system prompt 中的指令文本。要求：

1. **不重复原有标准**：不要复述上面已有的维度描述。你生成的指令是「追加/修改说明」，原有标准会完整保留在批改 prompt 中
2. **只调整教师提到的部分**：教师只说改哪个维度就改哪个，没提到的维度请勿提及或修改
3. **处理特殊情况**：如果教师标注了某学生的特殊原因（仪器故障、请假补做等），明确写出对该学生的处理方式
4. **保持严格与公正**：未被特殊标注的学生仍按原标准严格评分
5. **200-400 字，语言直接**：以「本次批改中，请特别注意以下调整：」开头，列出具体的修改点

输出格式：直接返回指令文本，不要加 markdown 代码块、标题或解释性前言。"""

    api_cfg = config.get("api", {})
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[{"role":"user","content":prompt}],
            temperature=0.7, max_tokens=2048,
            timeout=api_cfg.get("timeout",120))
        content = response.choices[0].message.content
        return content.strip() if content else ""
    except Exception as e:
        return f"（AI 优化失败：{e}）\n\n将使用教师的原始调整说明。\n\n{adjustment_notes}\n\n{special_situations}"

# ============================================================
# 6. Excel 生成（不变）
# ============================================================
def generate_excel(results, rubric, class_analysis, report_type):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = Path(tempfile.gettempdir()) / f"{report_type}_批改结果_{timestamp}.xlsx"
    wb = openpyxl.Workbook()
    hfont = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
    hfill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    halign = Alignment(horizontal="center", vertical="center", wrap_text=True)
    calign = Alignment(vertical="top", wrap_text=True)
    cc = Alignment(horizontal="center", vertical="center")
    border = Border(left=Side(style="thin"), right=Side(style="thin"),
                    top=Side(style="thin"), bottom=Side(style="thin"))
    dims = rubric["dimensions"]
    total_score = rubric["total_score"]

    # Sheet 1: 评分明细
    ws1 = wb.active
    ws1.title = "评分明细"
    headers = ["序号","学号","姓名","文件名"] + [f"{d['name']}\n({d['max_score']}分)" for d in dims] + [f"总分\n({total_score}分)","评语","状态"]
    for ci, h in enumerate(headers, 1):
        c = ws1.cell(row=1, column=ci, value=h); c.font = hfont; c.fill = hfill; c.alignment = halign; c.border = border
    for ri, r in enumerate(results, 2):
        si = r.get("student_info",{})
        ws1.cell(row=ri, column=1, value=ri-1).alignment = cc
        ws1.cell(row=ri, column=2, value=si.get("student_id","")).alignment = cc
        ws1.cell(row=ri, column=3, value=si.get("student_name","")).alignment = cc
        ws1.cell(row=ri, column=4, value=si.get("filename","")).alignment = calign
        if r.get("error"):
            for ci in range(5, 5+len(dims)): ws1.cell(row=ri, column=ci, value="N/A").alignment = cc
            ws1.cell(row=ri, column=5+len(dims), value="N/A").alignment = cc
            ws1.cell(row=ri, column=6+len(dims), value=r.get("message","未知错误")).alignment = calign
            ws1.cell(row=ri, column=7+len(dims), value="❌").alignment = cc
            ef = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            for ci in range(1, len(headers)+1): ws1.cell(row=ri, column=ci).fill = ef
        else:
            dm = {d["key"]:d for d in r.get("dimensions",[])}
            for di, dd in enumerate(dims):
                dv = dm.get(dd["key"],{})
                display = str(dv.get("score",0))
                if dv.get("comment"): display += f"\n{dv['comment']}"
                ws1.cell(row=ri, column=5+di, value=display).alignment = calign
            sc = ws1.cell(row=ri, column=5+len(dims), value=r.get("total_score",0)); sc.alignment = cc; sc.font = Font(bold=True, size=11)
            ws1.cell(row=ri, column=6+len(dims), value=r.get("overall_comment","")).alignment = calign
            ws1.cell(row=ri, column=7+len(dims), value="✅").alignment = cc
        for ci in range(1, len(headers)+1): ws1.cell(row=ri, column=ci).border = border
    widths = [5,16,10,30] + [28]*len(dims) + [10,50,8]
    for ci, w in enumerate(widths, 1): ws1.column_dimensions[get_column_letter(ci)].width = w
    ws1.freeze_panes = "A2"
    ws1.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(results)+1}"

    # Sheet 2: 班级统计
    ws2 = wb.create_sheet("班级统计")
    valid = [r for r in results if not r.get("error")]
    if valid:
        scores = [r["total_score"] for r in valid]
        avg_score = sum(scores)/len(scores)
        dist = {"优秀 (≥90%)":len([s for s in scores if s>=total_score*0.9]),
                "良好 (80-89%)":len([s for s in scores if total_score*0.8<=s<total_score*0.9]),
                "中等 (70-79%)":len([s for s in scores if total_score*0.7<=s<total_score*0.8]),
                "及格 (60-69%)":len([s for s in scores if total_score*0.6<=s<total_score*0.7]),
                "不及格 (<60%)":len([s for s in scores if s<total_score*0.6])}
        stats = [["指标","数值"],["报告类型",report_type],["提交总数",len(results)],["成功批改",len(valid)],
                 ["批改失败",len(results)-len(valid)],["满分",total_score],["最高分",round(max(scores),1)],
                 ["最低分",round(min(scores),1)],["平均分",round(avg_score,1)],["中位数",round(sorted(scores)[len(scores)//2],1)]]
        for ri,(l,v) in enumerate(stats,1):
            a=ws2.cell(row=ri,column=1,value=l); a.font=Font(name="微软雅黑",size=11,bold=True); a.alignment=cc
            b=ws2.cell(row=ri,column=2,value=v); b.alignment=cc
            if ri>=7: b.font=Font(size=12,bold=True,color="2F5496")
        ds=len(stats)+2; ws2.cell(row=ds,column=1,value="分数分布").font=Font(name="微软雅黑",size=12,bold=True)
        for i,(l,cnt) in enumerate(dist.items()):
            ws2.cell(row=ds+1+i,column=1,value=l).alignment=cc
            ws2.cell(row=ds+1+i,column=2,value=cnt).alignment=cc
            ws2.cell(row=ds+1+i,column=3,value="█"*cnt).font=Font(name="Consolas")
        dds=ds+len(dist)+3; ws2.cell(row=dds,column=1,value="各维度平均分").font=Font(name="微软雅黑",size=12,bold=True)
        for i,dim in enumerate(dims):
            dscores=[]
            for r in valid:
                for d in r.get("dimensions",[]):
                    if d["key"]==dim["key"]: dscores.append(d.get("score",0)); break
            da=sum(dscores)/len(dscores) if dscores else 0
            ws2.cell(row=dds+1+i,column=1,value=f"{dim['name']} (满分{dim['max_score']})").alignment=cc
            ws2.cell(row=dds+1+i,column=2,value=f"{round(da,2)}").alignment=cc
            pct=(da/dim["max_score"])*100 if dim["max_score"]>0 else 0
            ws2.cell(row=dds+1+i,column=3,value=f"{round(pct,0)}%").alignment=cc
        ws2.column_dimensions["A"].width=30; ws2.column_dimensions["B"].width=15; ws2.column_dimensions["C"].width=20
    else:
        ws2.cell(row=1,column=1,value="无有效评分数据").font=Font(color="FF0000")

    # Sheet 3: AI 班级分析
    ws3 = wb.create_sheet("AI班级分析")
    if class_analysis.get("summary"):
        ws3.cell(row=1,column=1,value="AI 班级整体分析").font=Font(name="微软雅黑",size=14,bold=True,color="2F5496")
        ws3.merge_cells("A1:D1")
        ws3.cell(row=3,column=1,value=class_analysis["summary"]).alignment=Alignment(wrap_text=True,vertical="top")
        ws3.column_dimensions["A"].width=80; ws3.column_dimensions["B"].width=20
        ws3.column_dimensions["C"].width=20; ws3.column_dimensions["D"].width=20; ws3.row_dimensions[3].height=300
    else:
        ws3.cell(row=1,column=1,value="未能生成班级分析").font=Font(color="FF0000")
    wb.save(output_path)
    return str(output_path)

# ============================================================
# 7. 班级 AI 分析（不变）
# ============================================================
def generate_class_analysis(client, model, results, rubric, experiment_name, config):
    valid = [r for r in results if not r.get("error")]
    if len(valid) < 3: return {"summary":"有效批改数量不足（需≥3份），无法进行班级分析。"}
    total_score = rubric["total_score"]; report_type = rubric["report_type"]
    scores = [r["total_score"] for r in valid]; avg_score = sum(scores)/len(scores)
    dim_summary = []
    for dim in rubric["dimensions"]:
        dscores=[]
        for r in valid:
            for d in r.get("dimensions",[]):
                if d["key"]==dim["key"]: dscores.append(d.get("score",0)); break
        da=sum(dscores)/len(dscores) if dscores else 0
        dim_summary.append(f"- {dim['name']}：平均 {round(da,2)}/{dim['max_score']} ({round(da/dim['max_score']*100,1)}%)")
    comments_summary = ""
    for r in valid[:10]:
        si=r.get("student_info",{})
        comments_summary+=f"\n{si.get('student_name','未知')} ({si.get('student_id','')})：{r.get('overall_comment','')[:200]}"
    prompt = f"""你是无机化学实验课的教师，正在查看一次{report_type}的批改汇总。
{f"实验名称：{experiment_name}" if experiment_name else ""}
全班情况：提交 {len(results)} 人，成功批改 {len(valid)} 人
平均分 {round(avg_score,2)}/{total_score}，最高 {max(scores)}，最低 {min(scores)}
各维度平均得分：{chr(10).join(dim_summary)}
部分评语摘录：{comments_summary[:3000]}
请做一个150-300字的班级整体分析，包括：整体水平判断、共性优点、共性不足、1-2条教学建议。直接返回分析文本。"""
    api_cfg = config.get("api", {})
    try:
        response = client.chat.completions.create(
            model=model, messages=[{"role":"user","content":prompt}],
            temperature=0.5, max_tokens=2048, timeout=api_cfg.get("timeout",120))
        return {"summary": response.choices[0].message.content.strip()}
    except Exception as e:
        return {"summary": f"生成班级分析失败：{str(e)}"}

# ============================================================
# 8. 文件查找 & 批量提取
# ============================================================
def find_report_files(zip_file, folder_path: str) -> list[str]:
    files = []
    if zip_file is not None:
        extract_dir = Path(tempfile.mkdtemp(prefix="reports_"))
        with zipfile.ZipFile(zip_file, "r") as zf: zf.extractall(extract_dir)
        for root, _, fnames in os.walk(extract_dir):
            for fn in fnames:
                if Path(fn).suffix.lower() in (".docx",".pdf") and not fn.startswith("~$"):
                    files.append(os.path.join(root, fn))
    elif folder_path and folder_path.strip():
        folder = Path(folder_path.strip())
        if folder.is_dir():
            for fp in folder.rglob("*"):
                if fp.suffix.lower() in (".docx",".pdf") and not fp.name.startswith("~$"):
                    files.append(str(fp))
        else: raise ValueError(f"文件夹不存在: {folder_path}")
    files.sort(key=lambda f: Path(f).name)
    return files

def load_all_files(zip_file, folder_path: str, with_images: bool = True) -> list[dict]:
    """提取所有文件的文本和图片，返回 [{filepath, filename, student_info, report_text, images}]"""
    filepaths = find_report_files(zip_file, folder_path)
    result = []
    for fp in filepaths:
        fn = Path(fp).name
        try:
            text = extract_text(fp)
        except Exception:
            text = ""
        si = extract_student_info(fn); si["filename"] = fn
        images = []
        if with_images:
            try:
                images = extract_images(fp)
            except Exception:
                pass
        findings = scan_report(fp)
        result.append({"filepath": fp, "filename": fn, "student_info": si, "report_text": text,
                       "images": images, "injection_findings": findings})
    return result

# ============================================================
# 9. 生成结果 HTML（复用）
# ============================================================
def build_result_html(result, rubric, show_filename=True):
    si = result.get("student_info", {})
    fn = si.get("filename", "")
    findings = result.get("injection_findings", [])
    security_html = ""
    if findings:
        rows = "".join(
            f"<li><b>第 {item.get('page', '?')} 页 · {html.escape(str(item.get('category', '可疑指令')))}</b>："
            f"“{html.escape(str(item.get('excerpt', '')))}”</li>"
            for item in findings
        )
        security_html = (
            "<div style='padding:10px;margin:6px 0;border:2px solid #d97706;background:#fff7ed;border-radius:6px'>"
            f"<b style='color:#b45309'>⚠️ 检测到可能影响 AI 评分的学生指令</b><br>"
            f"学生：{html.escape(str(si.get('student_name', '未识别')))}　"
            f"学号：{html.escape(str(si.get('student_id', '未识别')))}　"
            f"文件：{html.escape(str(fn))}<ul>{rows}</ul>"
            "<span style='color:#666'>这些文字已被标记为不可信内容，不会自动加扣分或认定违规，请教师核查。</span></div>"
        )
    if result.get("error"):
        return security_html + (f"<div style='padding:10px;margin:5px;border-left:3px solid red;background:#fff5f5'>"
                f"<b>❌ {si.get('student_name','?')} ({si.get('student_id','?')})</b>"
                f" — {fn}<br><span style='color:red'>{result.get('message','未知错误')}</span></div>")
    score = result.get("total_score", 0)
    max_score = rubric["total_score"]
    pct = score / max_score * 100 if max_score > 0 else 0
    if pct >= 90: sc = "#1a7a1a"
    elif pct >= 80: sc = "#2d8f2d"
    elif pct >= 70: sc = "#b8860b"
    elif pct >= 60: sc = "#cc6600"
    else: sc = "#cc0000"
    dh = ""
    for dim in result.get("dimensions", []):
        dd = next((d for d in rubric["dimensions"] if d["key"] == dim.get("key")), None)
        dn = dd["name"] if dd else dim.get("key","?")
        dm = dd["max_score"] if dd else 0
        dc = dim.get("comment","")
        dh += (f"<span style='display:inline-block;margin:3px;padding:4px 8px;"
               f"background:#f0f4ff;border-radius:4px;font-size:13px'>"
               f"<b>{dn}</b>：{dim.get('score',0)}/{dm}"
               f"{' <span style=color:#888>(' + dc[:60] + '...)</span>' if dc else ''}</span>")
    fn_html = f"<span style='font-size:11px;color:#aaa'>📄 {fn}</span>" if show_filename else ""
    return security_html + (f"<div style='padding:12px;margin:8px 0;border:1px solid #e0e0e0;border-radius:8px;background:#fafafa'>"
            f"<div style='display:flex;justify-content:space-between;align-items:center'>"
            f"<b style='font-size:16px'>{si.get('student_name','?')}</b>"
            f"<span style='color:#888;font-size:13px'>{si.get('student_id','?')}</span>"
            f"<span style='font-size:22px;font-weight:bold;color:{sc}'>{score}/{max_score}</span></div>"
            f"<div style='margin:8px 0'>{dh}</div>"
            f"<p style='color:#555;margin:5px 0;font-size:14px'>📝 {result.get('overall_comment','')}</p>"
            f"{fn_html}</div>")

def build_summary_html(results, rubric, class_analysis, label="📊 班级批改总结"):
    valid = [r for r in results if not r.get("error")]
    scores = [r["total_score"] for r in valid]
    avg_s = sum(scores)/len(scores) if scores else 0
    return (f"<div style='padding:12px;margin:15px 0;background:#f8fdf8;border:1px solid #b7d7b7;border-radius:8px'>"
            f"<h3 style='margin-top:0'>{label}</h3>"
            f"<p>共 <b>{len(results)}</b> 份，成功 <b>{len(valid)}</b> 份，"
            f"平均分 <b>{round(avg_s,2)}</b>/{rubric['total_score']}</p>"
            f"<p style='white-space:pre-wrap'>{class_analysis.get('summary','')}</p></div>")

# ============================================================
# 10. Gradio UI（重构：两轮批改工作流）
# ============================================================
def create_ui():
    config = load_config()
    api_defaults = config.get("api", {})
    default_base_url = api_defaults.get("base_url", "")

    with gr.Blocks(title="实验报告 AI 批改系统") as app:
        _theme = gr.themes.Soft(primary_hue="blue")
        _css = """
        .gradio-container { max-width: 1100px !important; }
        .results-scroll { max-height: 480px; overflow-y: auto; padding-right: 8px; border: 1px solid #e0e0e0; border-radius: 6px; padding: 10px; background: #fcfcfc; }
        """

        # ===== State =====
        file_data_state = gr.State([])       # [{filepath, filename, student_info, report_text}]
        first_results_state = gr.State([])   # 第一轮批改结果
        custom_instructions_state = gr.State("")  # 客制化指令文本
        final_results_state = gr.State([])   # 第二轮批改结果

        gr.Markdown("""# 🔬 实验报告 AI 批改系统
        **两轮批改流程**：① 快速预批改 → ② 查看结果 → ③ 调整评分规则 → ④ AI 优化规则 → ⑤ 重新批改 → ⑥ 下载 Excel""")

        # ===== Section 1: 配置栏 =====
        with gr.Row():
            api_key = gr.Textbox(label="API Key", placeholder="sk-xxxxxxxx", type="password", scale=2)
            base_url = gr.Textbox(label="API Base URL", value=default_base_url, scale=2)
            fetch_btn = gr.Button("🔄 拉取模型", size="sm", variant="secondary", scale=1)

        with gr.Row():
            model_dropdown = gr.Dropdown(label="选择模型", choices=api_defaults.get("fallback_models",[]),
                                         value=api_defaults.get("default_model",""), allow_custom_value=True, scale=2,
                                         info="可从下拉选，也可直接输入模型名")
            report_type = gr.Radio(choices=["预习报告","实验报告"], value="实验报告",
                                   label="报告类型", scale=1)
            experiment_name = gr.Textbox(label="实验名称（可选）", placeholder="如：三氯化六氨合钴(III)的制备和组分测定", scale=2)

        # ===== Section 2: 上传 & 预批改 =====
        gr.Markdown("### 📂 第一步：上传文件并快速预批改")
        with gr.Row():
            with gr.Column(scale=1):
                with gr.Tabs():
                    with gr.TabItem("📦 上传 ZIP"):
                        zip_upload = gr.File(label="选择 .zip 文件", file_types=[".zip"], type="filepath")
                    with gr.TabItem("📁 文件夹路径"):
                        folder_input = gr.Textbox(label="文件夹路径", placeholder="D:\\作业\\实验一\\")
                with gr.Row():
                    enable_images = gr.Checkbox(label="🖼️ 启用图片识别（需要视觉模型如 qwen3.5-27b）", value=True,
                                                info="提取报告中的图片（公式/图表/曲线），发送给多模态模型一起评分")
                pre_grade_btn = gr.Button("🚀 快速预批改", variant="primary", size="lg")

        # ===== Section 3: 第一轮结果 =====
        gr.Markdown("### 📊 预批改结果")
        with gr.Column(elem_classes="results-scroll"):
            first_results_html = gr.HTML(
                value="<p style='color:#888'>上传文件后点击「快速预批改」，AI 将用默认标准逐份批改</p>")

        # ===== Section 4: 调整面板（预批改后出现） =====
        with gr.Accordion("🔧 第二步：调整评分规则（预批改完成后展开）", open=False) as adjust_accordion:
            gr.Markdown("""
            根据预批改结果，你可以调整评分侧重或说明特殊情况。
            AI 会根据你的说明生成**客制化的评分规则**，然后重新批改所有报告。
            """)
            with gr.Row():
                with gr.Column(scale=1):
                    rubric_summary_md = gr.Markdown("*（预批改完成后自动显示当前评分标准）*")
                with gr.Column(scale=1):
                    adjustment_notes = gr.Textbox(
                        label="📝 评分侧重调整",
                        placeholder="例如：\n- 实验数据维度权重提升到35%\n- 原理维度放宽，只关注核心方程式\n- 本次重点关注有效数字的规范性",
                        lines=5, info="描述你希望调整的评分维度权重或关注重点")
                    special_situations = gr.Textbox(
                        label="⚠️ 特殊情况说明",
                        placeholder="例如：\n- 学生A的实验因仪器故障数据偏差，实验数据维度不扣分\n- 学生B的文献调研特别详实，格式维度+1分\n- 学生C因请假补做，操作分不扣",
                        lines=4, info="标注个别学生的特殊情况（加分/免扣分/豁免等）")

            with gr.Row():
                optimize_btn = gr.Button("🤖 AI 优化评分规则", variant="secondary", size="lg")
                re_grade_btn = gr.Button("🔄 用优化后的规则重新批改", variant="primary", size="lg", visible=False)

            optimized_rules_display = gr.Textbox(
                label="客制化评分规则（AI 生成，可手动修改）",
                placeholder="点击「AI 优化评分规则」后，这里会显示优化后的评分指令...",
                lines=8, interactive=True)

        # ===== Section 5: 最终结果 + 下载 =====
        gr.Markdown("### ✅ 最终批改结果")
        with gr.Column(elem_classes="results-scroll"):
            final_results_html = gr.HTML(
                value="<p style='color:#888'>完成规则优化后，点击「重新批改」获得最终结果</p>")
        with gr.Row():
            download_file = gr.File(label="📥 下载 Excel 汇总表", visible=False)

        # ===== 事件绑定 =====

        # 拉取模型
        def on_fetch(ak, bu):
            if not ak or not ak.strip(): return gr.Dropdown(choices=["请先输入 API Key"], value="请先输入 API Key")
            models = fetch_model_list(ak.strip(), get_effective_base_url(bu))
            if models: return gr.Dropdown(choices=models, value=models[0], interactive=True, allow_custom_value=True)
            fb = api_defaults.get("fallback_models",["qwen-plus","glm-4"])
            return gr.Dropdown(choices=fb, value=fb[0], interactive=True, allow_custom_value=True)
        fetch_btn.click(fn=on_fetch, inputs=[api_key, base_url], outputs=[model_dropdown])

        # 报告类型切换时更新评分标准摘要
        def update_rubric_summary(rt):
            cfg = load_config()
            rubric = cfg["preview_rubric"] if rt == "预习报告" else cfg["lab_rubric"]
            lines = [f"**{rubric['report_type']}**（满分 {rubric['total_score']} 分）\n"]
            for i, d in enumerate(rubric["dimensions"], 1):
                lines.append(f"{i}. **{d['name']}**：满分 {d['max_score']} 分，权重 {int(d['weight']*100)}%")
            return "\n".join(lines)
        report_type.change(fn=update_rubric_summary, inputs=[report_type], outputs=[rubric_summary_md])

        # === 快速预批改（Generator） ===
        def do_pre_grade(ak, bu, model, rt, zip_f, folder_p, exp_name, with_images):
            """Generator: 预批改所有报告"""
            if not ak or not ak.strip():
                yield "<p style='color:red'>❌ 请输入 API Key</p>", [], [], ""
                return
            cfg = load_config()
            base = get_effective_base_url(bu)
            rubric = cfg["preview_rubric"] if rt == "预习报告" else cfg["lab_rubric"]
            client = OpenAI(base_url=base, api_key=ak.strip())

            # 1. 加载文件
            try:
                files_data = load_all_files(zip_f, folder_p, with_images=with_images)
            except Exception as e:
                yield f"<p style='color:red'>❌ 文件加载失败：{e}</p>", [], [], ""
                return
            if not files_data:
                yield "<p style='color:red'>❌ 未找到任何 .docx 或 .pdf 文件</p>", [], [], ""
                return

            total = len(files_data)
            results = []
            html_parts = []

            # 2. 逐份批改
            for idx, fd in enumerate(files_data):
                img_info = f" (含{len(fd.get('images',[]))}张图)" if fd.get("images") else ""
                header = f"<p style='color:#2F5496;font-weight:bold'>⏳ 预批改进度：{idx+1}/{total} —— {fd['filename']}{img_info}</p>"
                if not fd["report_text"].strip():
                    r = {"error":True, "message":"文件为空或无法提取文本", "student_info":fd["student_info"]}
                    r["injection_findings"] = fd.get("injection_findings", [])
                    results.append(r)
                    html_parts.append(build_result_html(r, rubric))
                    yield header + "".join(html_parts), results, files_data, ""
                    continue
                grading = grade_one_report(client, model.strip(), fd["report_text"], fd["student_info"],
                                              rubric, exp_name, cfg, images=fd.get("images", []))
                grading["student_info"] = fd["student_info"]
                grading["injection_findings"] = fd.get("injection_findings", [])
                results.append(grading)
                html_parts.append(build_result_html(grading, rubric))
                yield header + "".join(html_parts), results, files_data, ""

            # 3. 班级分析
            yield "<p style='color:#2F5496;font-weight:bold'>📊 预批改完成，正在进行班级分析...</p>" + "".join(html_parts), results, files_data, ""
            class_analysis = generate_class_analysis(client, model.strip(), results, rubric, exp_name, cfg)
            summary = build_summary_html(results, rubric, class_analysis, "📊 预批改班级总结")
            final_html = summary + "".join(html_parts)
            yield final_html, results, files_data, ""

        pre_grade_btn.click(
            fn=do_pre_grade,
            inputs=[api_key, base_url, model_dropdown, report_type, zip_upload, folder_input, experiment_name, enable_images],
            outputs=[first_results_html, first_results_state, file_data_state, custom_instructions_state]
        )

        # === AI 优化评分规则 ===
        def do_optimize_rubric(ak, bu, model, rt, exp_name, adj_notes, special, first_results):
            if not ak or not ak.strip(): return "", "", gr.Button(visible=False)
            if not first_results: return "", "❌ 请先完成预批改", gr.Button(visible=False)
            cfg = load_config()
            base = get_effective_base_url(bu)
            rubric = cfg["preview_rubric"] if rt == "预习报告" else cfg["lab_rubric"]
            client = OpenAI(base_url=base, api_key=ak.strip())
            custom = generate_customized_rubric(client, model.strip(), rubric, adj_notes, special, first_results, exp_name, cfg)
            return custom, custom, gr.Button(visible=True)
        optimize_btn.click(
            fn=do_optimize_rubric,
            inputs=[api_key, base_url, model_dropdown, report_type, experiment_name,
                    adjustment_notes, special_situations, first_results_state],
            outputs=[optimized_rules_display, custom_instructions_state, re_grade_btn]
        )

        # === 重新批改（Generator） ===
        def do_re_grade(ak, bu, model, rt, exp_name, custom_instr, files_data):
            if not ak or not ak.strip():
                yield "<p style='color:red'>❌ 请输入 API Key</p>", None, []
                return
            if not files_data:
                yield "<p style='color:red'>❌ 请先完成预批改</p>", None, []
                return
            cfg = load_config()
            base = get_effective_base_url(bu)
            rubric = cfg["preview_rubric"] if rt == "预习报告" else cfg["lab_rubric"]
            report_type_name = rubric["report_type"]
            client = OpenAI(base_url=base, api_key=ak.strip())

            total = len(files_data)
            results = []
            html_parts = []

            for idx, fd in enumerate(files_data):
                img_info = f" (含{len(fd.get('images',[]))}张图)" if fd.get("images") else ""
                header = f"<p style='color:#E67E22;font-weight:bold'>🔄 二次批改进度：{idx+1}/{total} —— {fd['filename']}{img_info}</p>"
                if not fd["report_text"].strip():
                    r = {"error":True, "message":"文件为空", "student_info":fd["student_info"]}
                    r["injection_findings"] = fd.get("injection_findings", [])
                    results.append(r)
                    html_parts.append(build_result_html(r, rubric))
                    yield header + "".join(html_parts), None, results
                    continue
                grading = grade_one_report(client, model.strip(), fd["report_text"], fd["student_info"],
                                            rubric, exp_name, cfg, extra_instructions=custom_instr,
                                            images=fd.get("images", []))
                grading["student_info"] = fd["student_info"]
                grading["injection_findings"] = fd.get("injection_findings", [])
                results.append(grading)
                html_parts.append(build_result_html(grading, rubric, show_filename=False))
                yield header + "".join(html_parts), None, results

            # 班级分析 + Excel
            yield "<p style='color:#E67E22;font-weight:bold'>📊 正在进行班级分析并生成 Excel...</p>" + "".join(html_parts), None, results
            class_analysis = generate_class_analysis(client, model.strip(), results, rubric, exp_name, cfg)
            excel_path = generate_excel(results, rubric, class_analysis, report_type_name)
            summary = build_summary_html(results, rubric, class_analysis, "✅ 最终批改结果")
            final_html = summary + "".join(html_parts)
            yield final_html, gr.File(value=excel_path, visible=True), results

        re_grade_btn.click(
            fn=do_re_grade,
            inputs=[api_key, base_url, model_dropdown, report_type, experiment_name,
                    custom_instructions_state, file_data_state],
            outputs=[final_results_html, download_file, final_results_state]
        )

        # ===== 使用说明 =====
        with gr.Accordion("📖 使用说明", open=False):
            gr.Markdown("""
            ### 两轮批改工作流

            1. **配置** → 输入 API Key，拉取模型列表，选择报告类型
            2. **上传** → 上传 ZIP 压缩包或填写文件夹路径
            3. **快速预批改** → AI 用默认评分标准批改，结果实时显示在右侧
            4. **查看结果** → 浏览每份报告的分数和评语，了解班级整体情况
            5. **调整规则** → 展开「调整评分规则」面板：
               - 在「评分侧重调整」中描述希望改变的维度权重或关注点
               - 在「特殊情况说明」中标注个别学生的加分/免扣分情况
            6. **AI 优化规则** → 点击按钮，AI 根据你的调整生成客制化评分指令
               （可手动再修改 AI 生成的规则）
            7. **重新批改** → 用优化后的规则重新批改所有报告
            8. **下载 Excel** → 最终结果包含评分明细、班级统计、AI 分析三个 Sheet

            ### 关于评分标准
            - 预习报告（满分 100 分）：原理理解、步骤完整性、安全认知、预期预判、格式规范
            - 实验报告（满分 100 分）：数据质量、计算正确性、误差分析、现象分析、结论深度、格式规范
            - 在 `config.yaml` 中可永久修改评分标准
            """)

        gr.Markdown("<div style='text-align:center;color:#aaa;font-size:12px;margin-top:20px'>"
                     "Powered by AI · 本地运行，数据安全</div>")

    return app, _theme, _css

if __name__ == "__main__":
    app, theme, css = create_ui()
    app.launch(server_name="127.0.0.1", server_port=7868, share=False,
               inbrowser=True, theme=theme, css=css)
